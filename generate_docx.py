#!/usr/bin/env python3
"""
generate_docx.py

Creates a sample .docx file with headings and sample text.
"""
import sys

try:
    from docx import Document
except ImportError:
    sys.exit("Error: python-docx package is required. Install with 'pip install python-docx'.")

def main():
    doc = Document()

    # Add various heading levels with sample paragraphs
    doc.add_heading('Heading 1', level=1)
    doc.add_paragraph('This is a sample paragraph under Heading 1.')

    doc.add_heading('Heading 2', level=2)
    doc.add_paragraph('This is a sample paragraph under Heading 2.')

    doc.add_heading('Heading 3', level=3)
    doc.add_paragraph('This is a sample paragraph under Heading 3.')

    doc.add_heading('Heading 4', level=4)
    doc.add_paragraph('This is a sample paragraph under Heading 4.')

    # Add a bullet list
    doc.add_heading('Bullet List Example', level=2)
    bullet_items = ['First bullet point', 'Second bullet point', 'Third bullet point']
    for item in bullet_items:
        doc.add_paragraph(item, style='List Bullet')

    # Add a numbered list
    doc.add_heading('Numbered List Example', level=2)
    number_items = ['First numbered item', 'Second numbered item', 'Third numbered item']
    for item in number_items:
        doc.add_paragraph(item, style='List Number')

    # Add closing remarks
    doc.add_paragraph('Some closing remarks or additional sample text at the end of the document.')

    output_filename = 'sample.docx'
    doc.save(output_filename)
    print(f'Created {output_filename}')

if __name__ == '__main__':  # noqa: E402
    main()